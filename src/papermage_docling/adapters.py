"""
Document format adapters for converting between document representations.

This module provides adapter classes for converting between different document formats,
particularly focusing on DoclingDocument and PaperMage formats, with support for multiple
input sources (PDF, DOCX, HTML) and spatial relationship tracking.
"""

import logging
import os
from pathlib import Path
from typing import Dict, List, Any, Optional, Union, Tuple, Literal

# Import the actual DoclingDocument from docling_core
try:
    from docling_core.types import DoclingDocument
except ImportError:
    logging.warning("docling_core not found. DoclingDocument conversions will not be available.")
    # Define a stub class for type hints only
    class DoclingDocument:
        pass

# Import PaperMage document classes
from papermage_docling.converters.document import Document, Entity, Span, Box

# Import visualization capabilities
from papermage_docling.visualizers.pdf_visualizer import PdfVisualizer

# Optional imports for DOCX support
try:
    import docx
    from docx.document import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    logging.warning("python-docx not found. DOCX conversions will not be available.")
    DOCX_SUPPORT = False
    class DocxDocument:
        pass

# Optional imports for HTML support
try:
    from bs4 import BeautifulSoup
    HTML_SUPPORT = True
except ImportError:
    logging.warning("BeautifulSoup not found. HTML conversions will not be available.")
    HTML_SUPPORT = False

logger = logging.getLogger(__name__)


class FormatAdapter:
    """Base adapter class for format conversions with enhanced support for spatial relationships."""
    
    @staticmethod
    def to_target_format(doc):
        """Convert source document to target format."""
        raise NotImplementedError("Subclasses must implement to_target_format")
    
    @staticmethod
    def from_source_format(source_doc):
        """Convert from source format to destination format."""
        raise NotImplementedError("Subclasses must implement from_source_format")
    
    @staticmethod
    def visualize(doc, output_dir: Optional[Path] = None, interactive: bool = False):
        """Visualize the document (if supported)."""
        raise NotImplementedError("Subclasses must implement visualize if visualization is supported")


class PaperMageAdapter(FormatAdapter):
    """Adapter for converting between DoclingDocument and PaperMage format with enhanced spatial relationship tracking."""
    
    @staticmethod
    def to_papermage(doc: DoclingDocument) -> Document:
        """
        Convert DoclingDocument to PaperMage Document format.
        
        Args:
            doc: A DoclingDocument instance
            
        Returns:
            A PaperMage-compatible Document instance
        """
        logger.info("Converting DoclingDocument to PaperMage format")
        
        # Create a new Document
        papermage_doc = Document()
        
        # Initialize entity layers
        pages = []
        paragraphs = []
        sentences = []
        words = []
        tokens = []
        blocks = []
        figures = []
        tables = []
        
        # Extract full text and track position
        full_text = ""
        current_position = 0
        
        # Copy over document metadata
        papermage_doc.metadata = {
            "source": "docling",
            "version": "0.18.0",  # Match PaperMage v0.18+ as specified in PRD
        }
        
        if hasattr(doc, "metadata"):
            for key, value in doc.metadata.items():
                papermage_doc.metadata[key] = value
        
        # Process text items
        for text_item in doc.texts:
            item_start = current_position
            item_text = text_item.text
            
            # Add to full text
            full_text += item_text + "\n"
            current_position += len(item_text) + 1  # +1 for newline
            
            # Create bounding box
            item_box = Box(
                x0=text_item.bbox[0],
                y0=text_item.bbox[1],
                x1=text_item.bbox[2],
                y1=text_item.bbox[3],
                page=text_item.page - 1  # Convert from 1-indexed to 0-indexed
            )
            
            # Create text span
            item_span = Span(start=item_start, end=current_position - 1)
            
            # Add as paragraph and block
            paragraphs.append(Entity(spans=[item_span], boxes=[item_box], text=item_text))
            blocks.append(Entity(spans=[item_span], boxes=[item_box], text=item_text))
            
            # Process RTL metadata if present
            is_rtl = text_item.metadata.get("is_rtl", False)
            if is_rtl:
                blocks[-1].metadata["is_rtl"] = True
                paragraphs[-1].metadata["is_rtl"] = True
                if "rtl_marker" in text_item.metadata:
                    blocks[-1].metadata["rtl_marker"] = text_item.metadata["rtl_marker"]
                    paragraphs[-1].metadata["rtl_marker"] = text_item.metadata["rtl_marker"]
                if "rtl_segments" in text_item.metadata:
                    blocks[-1].metadata["rtl_segments"] = text_item.metadata["rtl_segments"]
                    paragraphs[-1].metadata["rtl_segments"] = text_item.metadata["rtl_segments"]
            
            # Extract sentences and words/tokens if available
            if "sentences" in text_item.metadata:
                for sentence_data in text_item.metadata["sentences"]:
                    sent_text = sentence_data.get("text", "")
                    sent_start = item_start + sentence_data.get("start_pos", 0)
                    sent_end = item_start + sentence_data.get("end_pos", len(sent_text))
                    
                    # Create sentence span and box
                    sent_span = Span(start=sent_start, end=sent_end)
                    sent_box = Box(
                        x0=sentence_data.get("bbox", [0, 0, 0, 0])[0],
                        y0=sentence_data.get("bbox", [0, 0, 0, 0])[1],
                        x1=sentence_data.get("bbox", [0, 0, 0, 0])[2],
                        y1=sentence_data.get("bbox", [0, 0, 0, 0])[3],
                        page=text_item.page - 1
                    )
                    
                    # Add sentence entity
                    sentences.append(Entity(spans=[sent_span], boxes=[sent_box], text=sent_text))
                    
                    # Process words if available
                    if "words" in sentence_data:
                        for word_data in sentence_data["words"]:
                            word_text = word_data.get("text", "")
                            word_start = item_start + word_data.get("start_pos", 0)
                            word_end = item_start + word_data.get("end_pos", len(word_text))
                            
                            # Create word span and box
                            word_span = Span(start=word_start, end=word_end)
                            word_box = Box(
                                x0=word_data.get("bbox", [0, 0, 0, 0])[0],
                                y0=word_data.get("bbox", [0, 0, 0, 0])[1],
                                x1=word_data.get("bbox", [0, 0, 0, 0])[2],
                                y1=word_data.get("bbox", [0, 0, 0, 0])[3],
                                page=text_item.page - 1
                            )
                            
                            # Add word entity
                            words.append(Entity(spans=[word_span], boxes=[word_box], text=word_text))
                            
                            # Add tokens (same as words for now)
                            tokens.append(Entity(spans=[word_span], boxes=[word_box], text=word_text))
        
        # Process table items
        for table_item in doc.tables:
            table_start = current_position
            
            # Extract table text representation (caption or placeholder)
            table_text = table_item.caption if table_item.caption else f"Table {table_item.id}"
            
            # Add to full text
            full_text += table_text + "\n"
            current_position += len(table_text) + 1  # +1 for newline
            
            # Create bounding box
            table_box = Box(
                x0=table_item.bbox[0],
                y0=table_item.bbox[1],
                x1=table_item.bbox[2],
                y1=table_item.bbox[3],
                page=table_item.page - 1  # Convert from 1-indexed to 0-indexed
            )
            
            # Create table span
            table_span = Span(start=table_start, end=current_position - 1)
            
            # Create table metadata
            table_metadata = {
                "num_rows": table_item.rows,
                "num_cols": table_item.columns,
                "cells": table_item.cells
            }
            
            # Add table metadata from original document
            for key, value in table_item.metadata.items():
                table_metadata[key] = value
            
            # Add table entity
            tables.append(Entity(
                spans=[table_span],
                boxes=[table_box],
                text=table_text,
                metadata=table_metadata
            ))
            
            # Tables are also blocks
            blocks.append(Entity(
                spans=[table_span],
                boxes=[table_box],
                text=table_text,
                metadata={"is_table": True}
            ))
        
        # Process figure items
        for figure_item in doc.figures:
            figure_start = current_position
            
            # Extract figure caption if available
            caption_text = figure_item.caption if figure_item.caption else f"Figure {figure_item.id}"
            
            # Add to full text
            full_text += caption_text + "\n"
            current_position += len(caption_text) + 1  # +1 for newline
            
            # Create bounding box
            figure_box = Box(
                x0=figure_item.bbox[0],
                y0=figure_item.bbox[1],
                x1=figure_item.bbox[2],
                y1=figure_item.bbox[3],
                page=figure_item.page - 1  # Convert from 1-indexed to 0-indexed
            )
            
            # Create figure span
            figure_span = Span(start=figure_start, end=current_position - 1)
            
            # Create figure metadata
            figure_metadata = {
                "has_caption": figure_item.caption is not None,
                "image_data": figure_item.image_data
            }
            
            # Add figure metadata from original document
            for key, value in figure_item.metadata.items():
                figure_metadata[key] = value
            
            # Add figure entity
            figures.append(Entity(
                spans=[figure_span],
                boxes=[figure_box],
                text=caption_text,
                metadata=figure_metadata
            ))
            
            # Figures are also blocks
            blocks.append(Entity(
                spans=[figure_span],
                boxes=[figure_box],
                text=caption_text,
                metadata={"is_figure": True}
            ))
        
        # Process key-value items
        for kv_item in doc.key_value_items:
            kv_start = current_position
            
            # Format as "key: value"
            kv_text = f"{kv_item.key}: {kv_item.value}"
            
            # Add to full text
            full_text += kv_text + "\n"
            current_position += len(kv_text) + 1  # +1 for newline
            
            # Create bounding box
            kv_box = Box(
                x0=kv_item.bbox[0],
                y0=kv_item.bbox[1],
                x1=kv_item.bbox[2],
                y1=kv_item.bbox[3],
                page=kv_item.page - 1  # Convert from 1-indexed to 0-indexed
            )
            
            # Create key-value span
            kv_span = Span(start=kv_start, end=current_position - 1)
            
            # Add as a block with special metadata
            blocks.append(Entity(
                spans=[kv_span],
                boxes=[kv_box],
                text=kv_text,
                metadata={"is_key_value": True, "key": kv_item.key, "value": kv_item.value}
            ))
        
        # Determine page structure based on items
        if doc.pages:
            # If we have explicit page information, use it
            num_pages = doc.pages
        else:
            # Otherwise infer from the highest page number
            page_numbers = []
            for item_list in [doc.texts, doc.tables, doc.figures, doc.key_value_items]:
                if item_list:
                    page_numbers.extend([item.page for item in item_list])
            
            num_pages = max(page_numbers) if page_numbers else 1
        
        # Create page entities (simplistic approach)
        pages_text = full_text.split("\n\n", num_pages - 1) if num_pages > 1 else [full_text]
        page_start = 0
        
        for page_idx, page_text in enumerate(pages_text):
            page_end = page_start + len(page_text)
            
            # Create page span and box
            page_span = Span(start=page_start, end=page_end)
            page_box = Box(
                x0=0,
                y0=0,
                x1=612,  # Letter size default in points
                y1=792,
                page=page_idx
            )
            
            # Add page entity
            pages.append(Entity(spans=[page_span], boxes=[page_box]))
            
            # Update for next page
            page_start = page_end + 1
        
        # Set full document text
        papermage_doc.symbols = full_text
        
        # Add entity layers to the document
        papermage_doc.add_entity_layer("pages", pages)
        papermage_doc.add_entity_layer("paragraphs", paragraphs)
        papermage_doc.add_entity_layer("blocks", blocks)
        
        # Only add non-empty layers
        if sentences:
            papermage_doc.add_entity_layer("sentences", sentences)
        if words:
            papermage_doc.add_entity_layer("words", words)
        if tokens:
            papermage_doc.add_entity_layer("tokens", tokens)
        if figures:
            papermage_doc.add_entity_layer("figures", figures)
        if tables:
            papermage_doc.add_entity_layer("tables", tables)
        
        # Add spatial relationship data
        papermage_doc.metadata["spatial_relationships"] = {
            "layout_mode": doc.metadata.get("layout_mode", "default"),
            "reading_order": doc.metadata.get("reading_order", "default"),
            "has_multiple_columns": doc.metadata.get("has_multiple_columns", False),
            "is_rtl_dominant": doc.metadata.get("is_rtl_dominant", False)
        }
        
        logger.info(f"Successfully converted DoclingDocument to PaperMage format with {len(papermage_doc.entities)} entity layers")
        return papermage_doc
    
    @staticmethod
    def from_papermage(pm_doc: Document) -> DoclingDocument:
        """
        Convert PaperMage Document to DoclingDocument format.
        
        Args:
            pm_doc: A PaperMage Document instance
            
        Returns:
            A DoclingDocument instance
        """
        logger.info("Converting PaperMage format to DoclingDocument")
        
        from docling_core.types import DoclingDocument, TextItem, TableItem, FigureItem, KeyValueItem
        import uuid
        
        # Create a new DoclingDocument
        doc = DoclingDocument(
            texts=[],
            tables=[],
            figures=[],
            key_value_items=[],
            metadata=pm_doc.metadata.copy() if hasattr(pm_doc, "metadata") else {}
        )
        
        # Add source information to metadata
        doc.metadata["source"] = "papermage"
        
        # Copy spatial relationship data if available
        if hasattr(pm_doc, "metadata") and "spatial_relationships" in pm_doc.metadata:
            for key, value in pm_doc.metadata["spatial_relationships"].items():
                doc.metadata[key] = value
        
        # Process PaperMage paragraphs as text items
        paragraphs = pm_doc.get_entity_layer("paragraphs") if hasattr(pm_doc, "get_entity_layer") else []
        for para in paragraphs:
            # Skip if no boxes
            if not para.boxes:
                continue
                
            box = para.boxes[0]  # Use first box
            text = para.text or ""
            
            # Create TextItem
            text_item = TextItem(
                id=str(uuid.uuid4()),
                text=text,
                page=box.page + 1,  # Convert from 0-indexed to 1-indexed
                bbox=[box.x0, box.y0, box.x1, box.y1],
                confidence=1.0,
                metadata={}
            )
            
            # Copy any RTL metadata
            if hasattr(para, "metadata"):
                if "is_rtl" in para.metadata:
                    text_item.metadata["is_rtl"] = para.metadata["is_rtl"]
                if "rtl_marker" in para.metadata:
                    text_item.metadata["rtl_marker"] = para.metadata["rtl_marker"]
                if "rtl_segments" in para.metadata:
                    text_item.metadata["rtl_segments"] = para.metadata["rtl_segments"]
            
            # Extract sentences and words if available
            sentences = pm_doc.get_entity_layer("sentences") if hasattr(pm_doc, "get_entity_layer") else []
            words = pm_doc.get_entity_layer("words") if hasattr(pm_doc, "get_entity_layer") else []
            
            # Find sentences contained in this paragraph
            para_sentences = []
            for sentence in sentences:
                # Check if sentence is contained in paragraph based on spans
                for para_span in para.spans:
                    for sent_span in sentence.spans:
                        if (sent_span.start >= para_span.start and sent_span.end <= para_span.end):
                            # Calculate relative position within paragraph
                            rel_start = sent_span.start - para_span.start
                            rel_end = sent_span.end - para_span.start
                            
                            # Create sentence data
                            sent_data = {
                                "text": sentence.text,
                                "start_pos": rel_start,
                                "end_pos": rel_end,
                                "bbox": [
                                    sentence.boxes[0].x0,
                                    sentence.boxes[0].y0,
                                    sentence.boxes[0].x1,
                                    sentence.boxes[0].y1
                                ] if sentence.boxes else [0, 0, 0, 0],
                                "words": []
                            }
                            
                            # Find words contained in this sentence
                            for word in words:
                                for word_span in word.spans:
                                    if (word_span.start >= sent_span.start and word_span.end <= sent_span.end):
                                        # Calculate relative position within sentence
                                        word_rel_start = word_span.start - sent_span.start
                                        word_rel_end = word_span.end - sent_span.start
                                        
                                        # Create word data
                                        word_data = {
                                            "text": word.text,
                                            "start_pos": word_rel_start,
                                            "end_pos": word_rel_end,
                                            "bbox": [
                                                word.boxes[0].x0,
                                                word.boxes[0].y0,
                                                word.boxes[0].x1,
                                                word.boxes[0].y1
                                            ] if word.boxes else [0, 0, 0, 0]
                                        }
                                        
                                        sent_data["words"].append(word_data)
                            
                            para_sentences.append(sent_data)
            
            # Add sentences to text item metadata if any were found
            if para_sentences:
                text_item.metadata["sentences"] = para_sentences
            
            doc.texts.append(text_item)
        
        # Process PaperMage tables
        tables = pm_doc.get_entity_layer("tables") if hasattr(pm_doc, "get_entity_layer") else []
        for table in tables:
            # Skip if no boxes
            if not table.boxes:
                continue
                
            box = table.boxes[0]  # Use first box
            
            # Extract table metadata
            rows = table.metadata.get("num_rows", 0) if hasattr(table, "metadata") else 0
            columns = table.metadata.get("num_cols", 0) if hasattr(table, "metadata") else 0
            cells = table.metadata.get("cells", []) if hasattr(table, "metadata") else []
            
            # Create TableItem
            table_item = TableItem(
                id=str(uuid.uuid4()),
                page=box.page + 1,  # Convert from 0-indexed to 1-indexed
                bbox=[box.x0, box.y0, box.x1, box.y1],
                rows=rows,
                columns=columns,
                cells=cells,
                caption=table.text if hasattr(table, "text") else None,
                metadata={}
            )
            
            # Copy additional metadata
            if hasattr(table, "metadata"):
                for key, value in table.metadata.items():
                    if key not in ["num_rows", "num_cols", "cells"]:
                        table_item.metadata[key] = value
            
            doc.tables.append(table_item)
        
        # Process PaperMage figures
        figures = pm_doc.get_entity_layer("figures") if hasattr(pm_doc, "get_entity_layer") else []
        for figure in figures:
            # Skip if no boxes
            if not figure.boxes:
                continue
                
            box = figure.boxes[0]  # Use first box
            
            # Create FigureItem
            figure_item = FigureItem(
                id=str(uuid.uuid4()),
                page=box.page + 1,  # Convert from 0-indexed to 1-indexed
                bbox=[box.x0, box.y0, box.x1, box.y1],
                caption=figure.text if hasattr(figure, "text") else None,
                image_data=figure.metadata.get("image_data") if hasattr(figure, "metadata") else None,
                metadata={}
            )
            
            # Copy additional metadata
            if hasattr(figure, "metadata"):
                for key, value in figure.metadata.items():
                    if key != "image_data":
                        figure_item.metadata[key] = value
            
            doc.figures.append(figure_item)
        
        # Process blocks to find key-value pairs
        blocks = pm_doc.get_entity_layer("blocks") if hasattr(pm_doc, "get_entity_layer") else []
        for block in blocks:
            # Only process blocks marked as key-value
            if not hasattr(block, "metadata") or not block.metadata.get("is_key_value", False):
                continue
                
            # Skip if no boxes
            if not block.boxes:
                continue
                
            box = block.boxes[0]  # Use first box
            
            # Extract key and value
            key = block.metadata.get("key", "")
            value = block.metadata.get("value", "")
            
            # Create KeyValueItem
            kv_item = KeyValueItem(
                id=str(uuid.uuid4()),
                key=key,
                value=value,
                page=box.page + 1,  # Convert from 0-indexed to 1-indexed
                bbox=[box.x0, box.y0, box.x1, box.y1],
                confidence=1.0,
                metadata={}
            )
            
            # Copy additional metadata
            if hasattr(block, "metadata"):
                for key, value in block.metadata.items():
                    if key not in ["is_key_value", "key", "value"]:
                        kv_item.metadata[key] = value
            
            doc.key_value_items.append(kv_item)
        
        logger.info(f"Successfully converted PaperMage document to DoclingDocument with {len(doc.texts)} texts, {len(doc.tables)} tables, {len(doc.figures)} figures")
        return doc
    
    @staticmethod
    def visualize(doc, output_dir: Optional[Path] = None, interactive: bool = False, category: Literal["all", "char", "word", "line"] = "all"):
        """
        Visualize a document using the PdfVisualizer.
        
        Args:
            doc: Document to visualize (PaperMage or DoclingDocument)
            output_dir: Directory to save visualizations
            interactive: Whether to display visualizations interactively
            category: Text unit category to visualize (char, word, line, all)
            
        Returns:
            None
        """
        # Import only when needed
        from papermage_docling.visualizers.pdf_visualizer import PdfVisualizer
        
        # Get the document source path
        if hasattr(doc, "metadata") and "source_path" in doc.metadata:
            pdf_path = doc.metadata["source_path"]
        else:
            logger.error("Cannot visualize document: no source_path in metadata")
            return
        
        # Create visualizer
        visualizer = PdfVisualizer(
            log_level="info",
            display_text=True,
            interactive=interactive,
            output_dir=output_dir
        )
        
        # Visualize document
        visualizer.visualize_pdf(
            pdf_path=pdf_path,
            category=category
        )
        
        logger.info(f"Document visualization completed. Output directory: {output_dir}")


class DocxAdapter(FormatAdapter):
    """Adapter for converting between DOCX and DoclingDocument/PaperMage formats."""
    
    @staticmethod
    def to_docling(docx_path: Union[str, Path]) -> DoclingDocument:
        """
        Convert a DOCX document to DoclingDocument format.
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            A DoclingDocument instance
        """
        if not DOCX_SUPPORT:
            raise ImportError("python-docx is required for DOCX conversion. Install with 'pip install python-docx'")
        
        logger.info(f"Converting DOCX document to DoclingDocument: {docx_path}")
        
        # Import required classes
        from docling_core.types import DoclingDocument, TextItem, TableItem, FigureItem, KeyValueItem
        import uuid
        
        # Load the DOCX document
        doc_obj = docx.Document(docx_path)
        
        # Create a new DoclingDocument
        doc = DoclingDocument(
            texts=[],
            tables=[],
            figures=[],
            key_value_items=[],
            metadata={
                "source": "docx",
                "source_path": str(docx_path),
                "title": os.path.basename(str(docx_path))
            }
        )
        
        # Set page counter (DOCX doesn't have direct page mapping)
        current_page = 1
        
        # Process paragraphs
        for i, para in enumerate(doc_obj.paragraphs):
            # Skip empty paragraphs
            if not para.text.strip():
                continue
            
            # Create simplified bounding box (DOCX doesn't provide spatial info)
            # Using default values that will be updated by renderers
            bbox = [0, i * 20, 500, (i + 1) * 20]
            
            # Create text item
            text_item = TextItem(
                id=str(uuid.uuid4()),
                text=para.text,
                page=current_page,
                bbox=bbox,
                confidence=1.0,
                metadata={
                    "style": para.style.name,
                    "is_heading": para.style.name.startswith("Heading"),
                    "heading_level": int(para.style.name.replace("Heading ", "")) 
                        if para.style.name.startswith("Heading") and para.style.name[8:].isdigit() 
                        else None
                }
            )
            
            doc.texts.append(text_item)
            
            # Increment page counter every 30 paragraphs (arbitrary)
            if i > 0 and i % 30 == 0:
                current_page += 1
        
        # Process tables
        for i, table in enumerate(doc_obj.tables):
            rows = len(table.rows)
            cols = len(table.columns)
            
            # Extract table content
            cells = []
            for r in range(rows):
                for c in range(cols):
                    try:
                        cell_text = table.cell(r, c).text
                        cells.append({
                            "row": r,
                            "col": c,
                            "text": cell_text,
                            "rowspan": 1,
                            "colspan": 1
                        })
                    except Exception as e:
                        logger.warning(f"Error extracting table cell ({r},{c}): {e}")
            
            # Create simplified bounding box
            bbox = [50, i * 100, 550, (i + 1) * 100]
            
            # Create table item
            table_item = TableItem(
                id=str(uuid.uuid4()),
                page=current_page,
                bbox=bbox,
                rows=rows,
                columns=cols,
                cells=cells,
                caption=f"Table {i+1}",
                metadata={}
            )
            
            doc.tables.append(table_item)
            
            # Increment page counter for each table
            current_page += 1
        
        # Set total pages
        doc.pages = current_page
        
        logger.info(f"Successfully converted DOCX to DoclingDocument with {len(doc.texts)} texts, {len(doc.tables)} tables")
        return doc
    
    @staticmethod
    def to_papermage(docx_path: Union[str, Path]) -> Document:
        """
        Convert a DOCX document to PaperMage Document format.
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            A PaperMage Document instance
        """
        # First convert to DoclingDocument, then to PaperMage
        docling_doc = DocxAdapter.to_docling(docx_path)
        
        # Then convert to PaperMage
        return PaperMageAdapter.to_papermage(docling_doc)


class HtmlAdapter(FormatAdapter):
    """Adapter for converting between HTML and DoclingDocument/PaperMage formats."""
    
    @staticmethod
    def to_docling(html_path: Union[str, Path]) -> DoclingDocument:
        """
        Convert an HTML document to DoclingDocument format.
        
        Args:
            html_path: Path to the HTML file
            
        Returns:
            A DoclingDocument instance
        """
        if not HTML_SUPPORT:
            raise ImportError("BeautifulSoup is required for HTML conversion. Install with 'pip install beautifulsoup4'")
        
        logger.info(f"Converting HTML document to DoclingDocument: {html_path}")
        
        # Import required classes
        from docling_core.types import DoclingDocument, TextItem, TableItem, KeyValueItem
        import uuid
        
        # Load and parse the HTML document
        with open(html_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Create a new DoclingDocument
        doc = DoclingDocument(
            texts=[],
            tables=[],
            figures=[],
            key_value_items=[],
            metadata={
                "source": "html",
                "source_path": str(html_path),
                "title": soup.title.text if soup.title else os.path.basename(str(html_path))
            }
        )
        
        # Extract metadata
        if soup.find('meta', attrs={'name': 'author'}):
            doc.metadata['author'] = soup.find('meta', attrs={'name': 'author'}).get('content', '')
        
        if soup.find('meta', attrs={'name': 'description'}):
            doc.metadata['description'] = soup.find('meta', attrs={'name': 'description'}).get('content', '')
        
        # Set page counter
        current_page = 1
        y_position = 0
        
        # Process headings and paragraphs
        for i, element in enumerate(soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div'])):
            # Skip empty elements
            if not element.text.strip():
                continue
            
            # Create simplified bounding box
            height = 20 * (len(element.text) // 80 + 1)  # Estimate height based on text length
            bbox = [50, y_position, 550, y_position + height]
            y_position += height + 10  # Add spacing
            
            # Create element metadata
            metadata = {}
            if element.name.startswith('h'):
                metadata['is_heading'] = True
                metadata['heading_level'] = int(element.name[1])
            
            # Extract element classes and IDs
            if element.get('class'):
                metadata['classes'] = element.get('class')
            if element.get('id'):
                metadata['id'] = element.get('id')
            
            # Create text item
            text_item = TextItem(
                id=str(uuid.uuid4()),
                text=element.text.strip(),
                page=current_page,
                bbox=bbox,
                confidence=1.0,
                metadata=metadata
            )
            
            doc.texts.append(text_item)
            
            # Increment page counter every 25 elements (arbitrary)
            if i > 0 and i % 25 == 0:
                current_page += 1
                y_position = 0
        
        # Process tables
        for i, table in enumerate(soup.find_all('table')):
            # Create simplified bounding box
            bbox = [50, i * 100, 550, (i + 1) * 100]
            
            rows = len(table.find_all('tr'))
            cols = max(len(row.find_all(['td', 'th'])) for row in table.find_all('tr'))
            
            # Extract table cells
            cells = []
            for r, row in enumerate(table.find_all('tr')):
                for c, cell in enumerate(row.find_all(['td', 'th'])):
                    # Extract rowspan and colspan
                    rowspan = int(cell.get('rowspan', 1))
                    colspan = int(cell.get('colspan', 1))
                    
                    cells.append({
                        "row": r,
                        "col": c,
                        "text": cell.text.strip(),
                        "rowspan": rowspan,
                        "colspan": colspan,
                        "is_header": cell.name == 'th'
                    })
            
            # Create table item
            table_item = TableItem(
                id=str(uuid.uuid4()),
                page=current_page,
                bbox=bbox,
                rows=rows,
                columns=cols,
                cells=cells,
                caption=f"Table {i+1}",
                metadata={}
            )
            
            doc.tables.append(table_item)
            
            # Increment page counter for each table
            current_page += 1
            y_position = 0
        
        # Extract key-value pairs from definition lists
        for i, dl in enumerate(soup.find_all('dl')):
            for dt, dd in zip(dl.find_all('dt'), dl.find_all('dd')):
                # Create simplified bounding box
                bbox = [50, y_position, 550, y_position + 20]
                y_position += 30  # Add spacing
                
                # Create key-value item
                kv_item = KeyValueItem(
                    id=str(uuid.uuid4()),
                    key=dt.text.strip(),
                    value=dd.text.strip(),
                    page=current_page,
                    bbox=bbox,
                    confidence=1.0,
                    metadata={}
                )
                
                doc.key_value_items.append(kv_item)
        
        # Set total pages
        doc.pages = current_page
        
        logger.info(f"Successfully converted HTML to DoclingDocument with {len(doc.texts)} texts, {len(doc.tables)} tables")
        return doc
    
    @staticmethod
    def to_papermage(html_path: Union[str, Path]) -> Document:
        """
        Convert an HTML document to PaperMage Document format.
        
        Args:
            html_path: Path to the HTML file
            
        Returns:
            A PaperMage Document instance
        """
        # First convert to DoclingDocument, then to PaperMage
        docling_doc = HtmlAdapter.to_docling(html_path)
        
        # Then convert to PaperMage
        return PaperMageAdapter.to_papermage(docling_doc)


# Factory for creating appropriate adapters based on file type
class AdapterFactory:
    """Factory for creating document format adapters based on file type."""
    
    @staticmethod
    def get_adapter_for_file(file_path: Union[str, Path]) -> FormatAdapter:
        """
        Get the appropriate adapter for the given file.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            A FormatAdapter instance
        """
        file_path_str = str(file_path).lower()
        
        if file_path_str.endswith('.pdf'):
            return PaperMageAdapter()
        elif file_path_str.endswith('.docx'):
            if not DOCX_SUPPORT:
                raise ImportError("python-docx is required for DOCX conversion.")
            return DocxAdapter()
        elif file_path_str.endswith(('.html', '.htm')):
            if not HTML_SUPPORT:
                raise ImportError("BeautifulSoup is required for HTML conversion.")
            return HtmlAdapter()
        else:
            raise ValueError(f"Unsupported file format: {file_path}") 